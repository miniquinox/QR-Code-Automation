# # Import docx NOT python-docx
# import os
# import docx
# from docx.shared import Mm

# # Create an instance of a word document
# doc = docx.Document()

# # Add a Title to the document
# doc.add_heading('GeeksForGeeks', 0)

# # Loop through all PNG files in directory
# directory = './'
# for filename in os.listdir(directory):
#     if filename.endswith(".png"):
#         # Image with defined size
#         doc.add_heading(f'Image: {filename}', 3)
#         doc.add_picture(os.path.join(directory, filename), width=Mm(27.8), height=Mm(27.8))

# # Now save the document to a location
# doc.save('qrritos.docx')


# Import docx NOT python-docx
import os
import docx
from docx.shared import Mm
from docx.enum.text import WD_LINE_SPACING

# Create an instance of a word document
doc = docx.Document()

# Set the line spacing of the whole document to 1.5
doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Add a Title to the document
doc.add_heading('QRRITOS AL PODER', 0)

# Create a paragraph to hold the images
images_paragraph = doc.add_paragraph()

# Loop through all PNG files in directory
directory = './'
for filename in os.listdir(directory):
    if filename.endswith(".png"):
        # Image with defined size
        images_paragraph.add_run().add_picture(os.path.join(directory, filename), width=Mm(27.8), height=Mm(27.8))
        # Add two spaces between each image
        images_paragraph.add_run("  ")

# Now save the document to a location
doc.save('qrritos.docx')

